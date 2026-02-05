"""
Generate AX3 Digital Transformation Discovery Questions Word Document
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_question(doc, number, question_text, options=None, context=None):
    """Add a numbered question with optional multiple-choice options and context."""
    # Question
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {question_text}")
    run.bold = True
    
    # Context if provided
    if context:
        context_p = doc.add_paragraph()
        context_run = context_p.add_run(f"Context: {context}")
        context_run.italic = True
        context_run.font.size = Pt(10)
    
    # Multiple choice options if provided
    if options:
        for opt in options:
            doc.add_paragraph(f"☐ {opt}", style='List Bullet')
        doc.add_paragraph("☐ Other: _________________________________")
    
    # Response field
    doc.add_paragraph("Your Response:")
    response_p = doc.add_paragraph("_" * 80)
    doc.add_paragraph()  # Spacing

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('AX3 Digital Transformation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Discovery Questions for Technical Specification', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run("Document Purpose").bold = True
    doc.add_paragraph(
        "This document contains questions that need clarification before detailed implementation "
        "can proceed. Your responses will help us finalize the technical specification and reduce "
        "project risk. Please complete this document at your convenience and return it to the project team."
    )
    
    doc.add_paragraph()
    instructions = doc.add_paragraph()
    instructions.add_run("Instructions").bold = True
    doc.add_paragraph(
        "• For multiple-choice questions, check (☐) all options that apply\n"
        "• For open-ended questions, provide as much detail as possible\n"
        "• If you're unsure about an answer, please indicate that and provide your best estimate\n"
        "• Feel free to add additional context or attach supporting documents"
    )
    
    doc.add_paragraph()
    doc.add_paragraph(f"Date Issued: February 5, 2026")
    doc.add_paragraph("Response Due: ________________")
    doc.add_paragraph("Respondent Name: ________________")
    doc.add_paragraph("Respondent Role: ________________")
    
    # =========================================================================
    # GROUP 1: AS400 & Legacy System Integration
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('Group 1: AS400 & Legacy System Integration', level=1)
    doc.add_paragraph(
        "These questions relate to creating minimal AS400 records for digital-first products "
        "and understanding AS400 capabilities for the new digital product path."
    )
    doc.add_paragraph()
    
    add_question(doc, 1,
        "What fields are absolutely required to create a minimal AS400 product record for sales tracking purposes?",
        options=[
            "Pub Number",
            "Item Number", 
            "Title",
            "Product Type Code",
            "Price",
            "Royalty Flag",
            "Status",
            "Territory"
        ],
        context="We need to understand the minimum viable record to track digital-first product sales."
    )
    
    add_question(doc, 2,
        "Can a new Product Type Code (e.g., 'D' for digital-native) be added to AS400?",
        options=[
            "Yes, this is a configuration change only",
            "Yes, but requires code changes",
            "No, we must use an existing product type",
            "Unsure - need to check with AS400 administrator"
        ],
        context="A new product type would help distinguish digital-first products from physical-derived ones."
    )
    
    add_question(doc, 3,
        "Who is the appropriate contact for AS400 configuration or code changes?",
        context="We need to coordinate any AS400 modifications."
    )
    
    add_question(doc, 4,
        "How does the AS400 product status field work? Which values indicate an active, sellable product?",
        options=[
            "CUR (Current)",
            "TOP",
            "NEW",
            "PND (Pending)",
            "Other: ________________"
        ],
        context="The validator needs to verify products have the correct status before upload."
    )
    
    add_question(doc, 5,
        "Is there existing documentation for the AS400 product record structure and field definitions?",
        options=[
            "Yes, available internally",
            "Yes, but outdated",
            "No formal documentation exists",
            "Unsure"
        ]
    )
    
    # =========================================================================
    # GROUP 2: Portal (portal.ax3.com) & Development
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('Group 2: Portal (portal.ax3.com) Technical Details', level=1)
    doc.add_paragraph(
        "These questions help us understand the portal's technical constraints and plan "
        "for minimal modifications or potential bypass strategies."
    )
    doc.add_paragraph()
    
    add_question(doc, 6,
        "What technology stack is portal.ax3.com built on?",
        options=[
            "ASP.NET Framework (specify version: ________)",
            "ASP.NET Core (specify version: ________)",
            "PHP",
            "Java",
            "Node.js",
            "Other: ________________"
        ],
        context="Understanding the technology helps us estimate modification effort."
    )
    
    add_question(doc, 7,
        "What is the deployment process for portal changes?",
        options=[
            "Direct deployment to production",
            "Staging environment testing first",
            "Formal change management process",
            "CI/CD pipeline",
            "Manual deployment by specific team member"
        ]
    )
    
    add_question(doc, 8,
        "Who has access to make portal code changes, and what is their availability?",
        context="We understand Lee may be involved. Please clarify his role and availability."
    )
    
    add_question(doc, 9,
        "What database tables does the portal write to when processing a digital product upload?",
        context="If portal modifications prove difficult, we may need to bypass the portal by writing directly to its target tables."
    )
    
    add_question(doc, 10,
        "Is there a staging or test environment for portal.ax3.com?",
        options=[
            "Yes, full staging environment",
            "Yes, but limited/outdated",
            "No staging environment",
            "Development environment only"
        ]
    )
    
    add_question(doc, 11,
        "Can the portal accept a null or empty ref_sku (reference SKU) field in the upload template?",
        options=[
            "Yes, already supported",
            "No, ref_sku is currently required",
            "Unsure - needs testing",
            "Would require code change"
        ],
        context="Digital-first products won't have a physical parent SKU."
    )
    
    # =========================================================================
    # GROUP 3: FileMaker & Data Sources
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('Group 3: FileMaker & Data Sources', level=1)
    doc.add_paragraph(
        "These questions clarify data sources and field mappings needed for automated template generation."
    )
    doc.add_paragraph()
    
    add_question(doc, 12,
        "Can we obtain the FileMaker field definition spreadsheets showing all fields and their mappings?",
        options=[
            "Yes, available immediately",
            "Yes, but needs to be compiled",
            "Partially available",
            "No, we would need to reverse-engineer from SQL"
        ],
        context="These spreadsheets will enable accurate template auto-generation."
    )
    
    add_question(doc, 13,
        "Which FileMaker tables are included in the nightly CSV export to SQL Server?",
        context="We want to confirm all needed data is already available in SQL Server."
    )
    
    add_question(doc, 14,
        "For Performance Music products, how is WebCRD accessed and what metadata does it contain?",
        options=[
            "WebCRD has a direct database connection",
            "WebCRD has an API",
            "WebCRD data is exported to files",
            "WebCRD data is manually entered",
            "Unsure of WebCRD integration method"
        ],
        context="We need to determine how to retrieve Performance Music metadata for template generation."
    )
    
    add_question(doc, 15,
        "What are the valid values for the OutputToWeb field in FileMaker?",
        options=[
            "JJ114",
            "DIGIONLY", 
            "DIGITALMASTER (proposed new value)",
            "Other: ________________"
        ],
        context="We may need to add DIGITALMASTER as a new value for digital-first products."
    )
    
    add_question(doc, 16,
        "Who manages FileMaker and can add new field values (like OutputToWeb options)?",
        context="We need to coordinate any FileMaker configuration changes."
    )
    
    # =========================================================================
    # GROUP 4: Template & Upload Process
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('Group 4: Template & Upload Process', level=1)
    doc.add_paragraph(
        "These questions clarify the current template structure and upload validation requirements."
    )
    doc.add_paragraph()
    
    add_question(doc, 17,
        "Are there documented specifications for each template type (Performance, Sheets, Choral)?",
        options=[
            "Yes, formal documentation exists",
            "Yes, but in the templates themselves (column headers)",
            "Informal/tribal knowledge only",
            "No documentation"
        ]
    )
    
    add_question(doc, 18,
        "Besides Column AK, are there other columns or fields that cause upload failures if present or formatted incorrectly?",
        context="We want to build comprehensive validation rules."
    )
    
    add_question(doc, 19,
        "What are the most common reasons for portal upload failures currently?",
        context="Understanding failure patterns helps us build better validation."
    )
    
    add_question(doc, 20,
        "For the _CC (Condensed Score) special case, are there any other rules beyond sort order = 1 and page count = 1?",
        options=[
            "No, those are the only special rules",
            "Yes, additional rules exist (please describe below)",
            "Unsure"
        ]
    )
    
    add_question(doc, 21,
        "How should UPC consistency be validated across set parts? What defines a 'set'?",
        context="We need to understand the business rules for UPC validation."
    )
    
    # =========================================================================
    # GROUP 5: MRID & Dealer Distribution
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('Group 5: MRID & Dealer Distribution', level=1)
    doc.add_paragraph(
        "These questions clarify the MRID file generation and FTP sync process."
    )
    doc.add_paragraph()
    
    add_question(doc, 22,
        "Can you provide a sample of the MR_Upload_Dealer_Full.txt and MR_Upload_Dealer.txt files?",
        options=[
            "Yes, sample files can be provided",
            "Yes, files are available on the FTP server",
            "No samples available"
        ],
        context="Sample files will help us understand the exact format required."
    )
    
    add_question(doc, 23,
        "What is the file format specification for MRID dealer files?",
        options=[
            "Tab-delimited",
            "Comma-delimited (CSV)",
            "Fixed-width",
            "XML",
            "Other: ________________"
        ]
    )
    
    add_question(doc, 24,
        "What triggers the need for a delta file (MR_Upload_Dealer.txt) vs. full file?",
        options=[
            "Delta generated on every change",
            "Delta generated daily",
            "Delta generated on request",
            "Both files generated together",
            "Other schedule: ________________"
        ]
    )
    
    add_question(doc, 25,
        "Are there any dealers or partners who should NOT receive certain digital products?",
        context="This helps us understand if the ax3.com-only flag needs additional distribution rules."
    )
    
    # =========================================================================
    # GROUP 6: ax3.com Website & Display
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('Group 6: ax3.com Website & Display', level=1)
    doc.add_paragraph(
        "These questions help us understand how digital products appear on the customer-facing website."
    )
    doc.add_paragraph()
    
    add_question(doc, 26,
        "What data source does ax3.com use to display product information?",
        options=[
            "Direct SQL Server query",
            "Cached/replicated data",
            "API calls to another service",
            "Combination of sources",
            "Unsure"
        ],
        context="We need to verify digital-first products will display correctly."
    )
    
    add_question(doc, 27,
        "Are there known dependencies in ax3.com's product display on having a physical ref_sku?",
        options=[
            "Yes, ref_sku is used in display logic",
            "No, ref_sku is not used for display",
            "Unsure - needs testing",
            "Varies by product type"
        ],
        context="Digital-first products won't have a physical parent reference."
    )
    
    add_question(doc, 28,
        "Who manages ax3.com and can make display-related fixes if needed?",
    )
    
    add_question(doc, 29,
        "Is there a test/staging version of ax3.com where we can validate digital-first product display?",
        options=[
            "Yes, full staging site",
            "Yes, but limited functionality",
            "No staging environment",
            "Can test with unlisted products on production"
        ]
    )
    
    # =========================================================================
    # GROUP 7: Royalty System Integration
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('Group 7: Royalty System Integration', level=1)
    doc.add_paragraph(
        "These questions clarify how digital products integrate with royalty tracking."
    )
    doc.add_paragraph()
    
    add_question(doc, 30,
        "How does the royalty system currently receive digital product information?",
        options=[
            "Via AS400 nightly sync",
            "Direct database connection",
            "Manual entry",
            "File-based import",
            "Other: ________________"
        ]
    )
    
    add_question(doc, 31,
        "Can the royalty system accept digital rights/royalties without a physical product parent?",
        options=[
            "Yes, already supported",
            "No, physical parent required",
            "Unsure - needs verification",
            "Would require system changes"
        ],
        context="Digital-first products need royalty tracking without physical counterparts."
    )
    
    add_question(doc, 32,
        "Who is the appropriate contact for royalty system questions and potential modifications?",
    )
    
    # =========================================================================
    # GROUP 8: Infrastructure & Access
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('Group 8: Infrastructure & Access', level=1)
    doc.add_paragraph(
        "These questions confirm access requirements and infrastructure availability."
    )
    doc.add_paragraph()
    
    add_question(doc, 33,
        "What hosting options are available for new services (Orchestration Service, Digital Product API)?",
        options=[
            "On-premises VMs",
            "Azure cloud services",
            "AWS cloud services",
            "Container hosting (Docker/Kubernetes)",
            "Combination available"
        ]
    )
    
    add_question(doc, 34,
        "What is the preferred technology stack for new API development?",
        options=[
            "ASP.NET Core (C#)",
            "Python (FastAPI/Flask)",
            "Node.js",
            "Java",
            "No preference - developer's choice"
        ]
    )
    
    add_question(doc, 35,
        "Please confirm access availability for the following systems (check all that can be provided):",
        options=[
            "SQL Server (read/write to new tables)",
            "AWS S3 credentials (alfred-dsm-pdfs, alfred-catfiles buckets)",
            "Dropbox API token",
            "FTP credentials (ftp1.ax3.com)",
            "Report Server access (alfredawssql06)"
        ]
    )
    
    add_question(doc, 36,
        "Are there any security, compliance, or approval requirements for creating new database tables or services?",
        options=[
            "No special approvals needed",
            "IT security review required",
            "Change management board approval",
            "Compliance review needed",
            "Other: ________________"
        ]
    )
    
    # =========================================================================
    # Additional Comments Section
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('Additional Information', level=1)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Please provide any additional context, concerns, or information that would help with the technical implementation:").bold = True
    
    doc.add_paragraph()
    for _ in range(15):
        doc.add_paragraph("_" * 85)
    
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Attachments:").bold = True
    doc.add_paragraph("Please attach any relevant documents (field mappings, sample files, system documentation):")
    doc.add_paragraph()
    doc.add_paragraph("☐ FileMaker field spreadsheets")
    doc.add_paragraph("☐ Sample MRID files")
    doc.add_paragraph("☐ Portal documentation")
    doc.add_paragraph("☐ AS400 record specifications")
    doc.add_paragraph("☐ Other: ________________")
    
    # Save document
    output_path = r"c:\Users\bened\Downloads\PKS\Music\future\AX3_Discovery_Questions.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    main()
